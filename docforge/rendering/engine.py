"""RenderingEngine — applies theme + layout decisions to a SemanticModel via python-docx."""

from __future__ import annotations

from pathlib import Path
from typing import Any

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from docforge.core.document import (
    Chapter,
    Heading,
    ImagePlaceholder,
    Paragraph,
    SemanticModel,
    Table,
)
from docforge.core.rendering import ChapterStyle, PageBalance, RenderingDecision
from docforge.logging.setup import get_logger

logger = get_logger(__name__)


# ---------------------------------------------------------------------------
# Typography tables: variant → (body_size_delta, line_spacing, spacing_before, spacing_after, body_font_override)
# ---------------------------------------------------------------------------
_TYPOGRAPHY: dict[str, tuple[float, float, int, int, str | None]] = {
    "conservative": (0,    1.15, 6,  6,  None),
    "editorial":    (0.5,  1.5,  8,  8,  None),
    "magazine":     (0,    1.15, 4,  4,  "Calibri"),
    "luxury":       (1,    1.35, 10, 10, None),
}

_PAGE_BALANCE_SPACING: dict[str, tuple[int, int]] = {
    "tight":    (3,  3),
    "balanced": (6,  6),
    "spacious": (12, 12),
}


def _hex_to_rgb(hex_colour: str) -> RGBColor:
    h = hex_colour.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


class RenderingEngine:
    def __init__(self, theme: dict) -> None:
        self._theme = theme

    def render(
        self,
        model: SemanticModel,
        decisions: dict[str, RenderingDecision],
        output_path: Path,
        language: str = "en",
        images: dict[int, Path] | None = None,
    ) -> Any:
        doc = docx.Document()
        self._apply_page_setup(doc)
        self._add_cover(doc, model, language)
        self._add_toc(doc, model, language)

        image_map = images or {}
        for chapter in model.chapters:
            decision = decisions.get(chapter.id)
            self._render_chapter(doc, chapter, decision, image_map)

        self._add_image_sources_appendix(doc, model, language, image_map)
        self._add_headers_footers(doc, language)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info("rendering_complete", output=str(output_path), images_inserted=len(image_map))
        return doc

    # ------------------------------------------------------------------
    # Page / section setup
    # ------------------------------------------------------------------

    def _apply_page_setup(self, doc: Any) -> None:
        spacing = self._theme.get("spacing", {})
        for section in doc.sections:
            section.top_margin    = Cm(spacing.get("page_margin_top_cm",    2.54))
            section.bottom_margin = Cm(spacing.get("page_margin_bottom_cm", 2.54))
            section.left_margin   = Cm(spacing.get("page_margin_left_cm",   2.54))
            section.right_margin  = Cm(spacing.get("page_margin_right_cm",  2.54))

    # ------------------------------------------------------------------
    # Cover + TOC
    # ------------------------------------------------------------------

    def _add_cover(self, doc: Any, model: SemanticModel, language: str) -> None:
        palette  = self._theme.get("palette", {})
        title    = model.chapters[0].title if model.chapters else "Document"

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(28)
        colour_hex = palette.get("primary", "#1a1a2e")
        run.font.color.rgb = _hex_to_rgb(colour_hex)
        doc.add_page_break()

    def _add_toc(self, doc: Any, model: SemanticModel, language: str) -> None:
        from docforge.core.i18n import get_label
        toc_label = get_label("toc", language)
        doc.add_heading(toc_label, level=1)
        for chapter in model.chapters:
            para = doc.add_paragraph(chapter.title, style="List Bullet")
            para.paragraph_format.left_indent = Cm(0)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # Chapter rendering
    # ------------------------------------------------------------------

    def _render_chapter(
        self,
        doc: Any,
        chapter: Chapter,
        decision: RenderingDecision | None,
        image_map: dict[int, Path],
    ) -> None:
        palette   = self._theme.get("palette", {})
        typography = self._theme.get("typography", {})

        chapter_style      = decision.chapter_style      if decision else ChapterStyle.STANDARD
        typography_variant = decision.typography_variant if decision else "conservative"
        page_balance       = decision.page_balance       if decision else PageBalance.BALANCED
        heading_colour     = decision.heading_colour     if decision else None
        pull_quote_flag    = decision.pull_quote         if decision else False
        sidebar_decision   = decision.sidebar            if decision else None

        # Resolved typography params
        ty = _TYPOGRAPHY.get(str(typography_variant), _TYPOGRAPHY["conservative"])
        body_size_delta, line_spacing, sp_before, sp_after, font_override = ty

        balance_sp = _PAGE_BALANCE_SPACING.get(str(page_balance), (6, 6))
        sp_before = balance_sp[0]
        sp_after  = balance_sp[1]

        base_body_size = typography.get("body_size", 11)
        body_size = base_body_size + body_size_delta
        body_font = font_override or typography.get("body_font", "Times New Roman")
        heading_font = typography.get("heading_font", "Times New Roman")

        # ------ Chapter heading ------
        heading_para = doc.add_heading("", level=chapter.heading_level)
        run = heading_para.add_run(chapter.title)
        run.font.name = heading_font

        # chapter_style drives heading size and decoration
        if chapter_style == ChapterStyle.OPENER:
            run.font.size = Pt(32)
            run.bold = True
            colour = heading_colour or palette.get("primary", "#1a1a2e")
            run.font.color.rgb = _hex_to_rgb(colour)
            self._add_rule(doc, colour)
        elif chapter_style == ChapterStyle.FEATURE:
            run.font.size = Pt(20)
            run.bold = True
            colour = heading_colour or palette.get("accent", "#0f3460")
            run.font.color.rgb = _hex_to_rgb(colour)
        else:  # standard
            run.font.size = Pt(typography.get("heading_size_h1", 16))
            if heading_colour:
                run.font.color.rgb = _hex_to_rgb(heading_colour)

        # ------ Sidebar (prepend a side-note before body) ------
        if sidebar_decision and sidebar_decision.enabled:
            first_para_text = next(
                (e.text for e in chapter.elements if isinstance(e, Paragraph) and e.text.strip()),
                None,
            )
            if first_para_text:
                self._add_sidebar(doc, first_para_text, sidebar_decision.type, palette)

        # ------ Pull quote from first long paragraph ------
        pull_quote_added = False

        for element in chapter.elements:
            if isinstance(element, Heading):
                h = doc.add_heading(element.text, level=element.level)
                h.runs[0].font.name = heading_font if h.runs else heading_font

            elif isinstance(element, Paragraph):
                if not element.text.strip():
                    continue

                # Insert pull quote once for the first paragraph with ≥ 40 words
                if (
                    pull_quote_flag
                    and not pull_quote_added
                    and len(element.text.split()) >= 40
                ):
                    self._add_pull_quote(doc, element.text, palette, body_font)
                    pull_quote_added = True

                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(sp_before)
                p.paragraph_format.space_after  = Pt(sp_after)
                p.paragraph_format.line_spacing = line_spacing

                run = p.add_run(element.text)
                run.font.name = body_font
                run.font.size = Pt(body_size)

                if element.style:
                    import contextlib
                    with contextlib.suppress(KeyError):
                        p.style = doc.styles[element.style]

            elif isinstance(element, Table):
                self._render_table(doc, element)

            elif isinstance(element, ImagePlaceholder):
                self._render_image_placeholder(doc, element, image_map)

    # ------------------------------------------------------------------
    # Decorative helpers
    # ------------------------------------------------------------------

    def _add_rule(self, doc: Any, colour_hex: str) -> None:
        """Thin horizontal rule paragraph (via bottom border on empty paragraph)."""
        from docx.oxml import OxmlElement
        para = doc.add_paragraph()
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), colour_hex.lstrip("#"))
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_pull_quote(self, doc: Any, text: str, palette: dict, body_font: str) -> None:
        """Blockquote-style pull quote — first sentence, large italic, indented."""
        sentence = text.split(".")[0].strip()
        if not sentence:
            return
        para = doc.add_paragraph()
        para.paragraph_format.left_indent  = Cm(1.5)
        para.paragraph_format.right_indent = Cm(1.5)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after  = Pt(12)
        run = para.add_run(f"“{sentence}.”")
        run.italic = True
        run.font.size = Pt(14)
        run.font.name = body_font
        accent = palette.get("accent", "#0f3460")
        run.font.color.rgb = _hex_to_rgb(accent)

    def _add_sidebar(self, doc: Any, text: str, sidebar_type: str, palette: dict) -> None:
        """Sidebar as a single-row table: narrow right column with a coloured background hint."""
        # Truncate to first sentence / 120 chars
        snippet = (text.split(".")[0].strip() + ".")[:120]
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        left_cell  = table.cell(0, 0)
        right_cell = table.cell(0, 1)

        # Make right cell narrow (sidebar)
        right_cell.width = Cm(4)
        left_cell.paragraphs[0].text = ""   # filled by caller's body elements
        sp = right_cell.paragraphs[0]
        run = sp.add_run(snippet)
        run.italic = True
        run.font.size = Pt(9)

        # Light background on sidebar cell
        from docx.oxml import OxmlElement
        tc_pr = right_cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F0F4F8")
        tc_pr.append(shd)
        doc.add_paragraph()

    # ------------------------------------------------------------------
    # Table rendering
    # ------------------------------------------------------------------

    def _render_table(self, doc: Any, table: Table) -> None:
        if not table.rows:
            return
        col_count = max(len(row.cells) for row in table.rows)
        doc_table = doc.add_table(rows=len(table.rows), cols=col_count)
        doc_table.style = "Table Grid"
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if c_idx < col_count:
                    text = " ".join(p.text for p in cell.content)
                    doc_table.cell(r_idx, c_idx).text = text

    # ------------------------------------------------------------------
    # Image placeholder rendering
    # ------------------------------------------------------------------

    def _render_image_placeholder(
        self,
        doc: Any,
        placeholder: ImagePlaceholder,
        image_map: dict[int, Path],
    ) -> None:
        img_path = image_map.get(id(placeholder))
        if img_path and img_path.exists():
            try:
                doc.add_picture(str(img_path), width=Inches(5.5))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logger.debug("image_inserted", path=str(img_path))
                return
            except Exception as exc:
                logger.warning("image_insert_failed", path=str(img_path), error=str(exc))
        p = doc.add_paragraph(f"[Image: {placeholder.placeholder_text.split(chr(10))[0]}]")
        p.runs[0].italic = True

    # ------------------------------------------------------------------
    # Appendix + headers/footers
    # ------------------------------------------------------------------

    def _add_image_sources_appendix(
        self,
        doc: Any,
        model: SemanticModel,
        language: str,
        image_map: dict[int, Path],
    ) -> None:
        from docforge.core.i18n import get_label
        appendix_label = get_label("image_sources", language)
        doc.add_page_break()
        doc.add_heading(appendix_label, level=1)
        if not image_map:
            doc.add_paragraph("No images were sourced in this rendering.")
        else:
            doc.add_paragraph(
                f"{len(image_map)} image(s) sourced via Wikimedia Commons (CC/PD licences)."
            )

    def _add_headers_footers(self, doc: Any, language: str) -> None:
        for section in doc.sections:
            header = section.header
            if not header.paragraphs:
                header.add_paragraph()
            header.paragraphs[0].text = ""

            footer = section.footer
            if not footer.paragraphs:
                footer.add_paragraph()
            para = footer.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not para.runs:
                para.add_run("")
