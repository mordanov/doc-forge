"""DocumentLoader — opens a .docx and builds the Document domain model.

The source file is NEVER modified (FR-002). All writes go to a copy.
"""

from __future__ import annotations

import uuid
from pathlib import Path
from typing import Any

import docx
from docx.oxml.ns import qn

from docforge.core.document import Document, DocumentMeta


class DocumentLoadError(Exception):
    pass


def load_document(path: Path) -> Document:
    """Load a .docx file and return an immutable Document domain model."""
    if not path.exists():
        raise DocumentLoadError(f"File not found: {path}")
    if path.suffix.lower() != ".docx":
        raise DocumentLoadError(f"Expected .docx, got: {path.suffix}")

    try:
        doc = docx.Document(str(path))
    except Exception as exc:
        raise DocumentLoadError(f"Cannot open {path}: {exc}") from exc

    core = doc.core_properties
    meta = DocumentMeta(
        author=core.author or None,
        subject=core.subject or None,
        keywords=[k.strip() for k in (core.keywords or "").split(",") if k.strip()],
        created=core.created,
        modified=core.modified,
    )

    title = _extract_title(doc, path)

    return Document(
        id=str(uuid.uuid4()),
        source_path=path.resolve(),
        title=title,
        language=_detect_language(doc),
        metadata=meta,
        sections=[],
    )


def _extract_title(doc: Any, path: Path) -> str:
    if doc.core_properties.title:
        return str(doc.core_properties.title)
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith("Heading 1") and para.text.strip():
            return str(para.text).strip()
    return path.stem


def _detect_language(doc: Any) -> str:
    try:
        body = doc.element.body
        for para in body.iter():
            lang_elem = para.find(qn("w:lang"))
            if lang_elem is not None:
                val = lang_elem.get(qn("w:val")) or ""
                if val:
                    return val.split("-")[0].lower()
    except Exception:
        pass
    return "en"
