"""Rendering tests — full pipeline against sample fixture."""

import tempfile
from pathlib import Path

import docx as docxlib
import pytest

SAMPLE_GUIDE = Path(__file__).parents[2] / "examples" / "sample-guide.docx"


@pytest.mark.skipif(not SAMPLE_GUIDE.exists(), reason="Sample fixture not found")
def test_render_pipeline_produces_valid_docx():
    from docforge import Renderer

    with tempfile.TemporaryDirectory() as tmp:
        output = Path(tmp) / "test_out.docx"
        report = Renderer().template("minimal").language("en").render(SAMPLE_GUIDE, output)

        assert report.succeeded(), f"Pipeline failed: {report.fatal_failure}"
        assert output.exists()
        assert output.stat().st_size > 0

        doc = docxlib.Document(str(output))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Table of Contents" in headings
        assert "Image Sources" in headings
        assert len(doc.paragraphs) > 5


@pytest.mark.skipif(not SAMPLE_GUIDE.exists(), reason="Sample fixture not found")
def test_render_pipeline_no_ai_flag():
    from docforge import Renderer

    with tempfile.TemporaryDirectory() as tmp:
        output = Path(tmp) / "test_noai.docx"
        report = Renderer().template("minimal").language("en").render(SAMPLE_GUIDE, output)
        assert report.succeeded()
        assert output.exists()


@pytest.mark.skipif(not SAMPLE_GUIDE.exists(), reason="Sample fixture not found")
def test_source_file_not_modified():
    import hashlib

    from docforge import Renderer

    checksum_before = hashlib.sha256(SAMPLE_GUIDE.read_bytes()).hexdigest()

    with tempfile.TemporaryDirectory() as tmp:
        output = Path(tmp) / "test_out.docx"
        Renderer().render(SAMPLE_GUIDE, output)

    checksum_after = hashlib.sha256(SAMPLE_GUIDE.read_bytes()).hexdigest()
    assert checksum_before == checksum_after, "Source file was modified!"
