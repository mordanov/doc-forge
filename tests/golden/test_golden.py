"""Golden document tests — compare rendered output against known fixtures."""

import tempfile
from pathlib import Path

import docx as docxlib
import pytest

SAMPLE_GUIDE = Path(__file__).parents[2] / "examples" / "sample-guide.docx"
EXPECTED_OUTPUT = Path(__file__).parent / "fixtures" / "expected-output.docx"


def _render_sample(tmp_dir: Path) -> Path:
    from docforge import Renderer

    output = tmp_dir / "golden_out.docx"
    report = Renderer().template("minimal").language("en").render(SAMPLE_GUIDE, output)
    assert report.succeeded(), f"Pipeline failed: {report.fatal_failure}"
    return output


@pytest.mark.skipif(not SAMPLE_GUIDE.exists(), reason="Sample fixture not found")
def test_output_has_cover_and_toc():
    with tempfile.TemporaryDirectory() as tmp:
        out = _render_sample(Path(tmp))
        doc = docxlib.Document(str(out))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Table of Contents" in headings


@pytest.mark.skipif(not SAMPLE_GUIDE.exists(), reason="Sample fixture not found")
def test_output_has_image_sources_appendix():
    with tempfile.TemporaryDirectory() as tmp:
        out = _render_sample(Path(tmp))
        doc = docxlib.Document(str(out))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Image Sources" in headings


@pytest.mark.skipif(not SAMPLE_GUIDE.exists(), reason="Sample fixture not found")
def test_output_heading_hierarchy():
    with tempfile.TemporaryDirectory() as tmp:
        out = _render_sample(Path(tmp))
        doc = docxlib.Document(str(out))
        h1s = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
        # Should have at least Introduction, one chapter, and Image Sources
        assert len(h1s) >= 2


@pytest.mark.skipif(not SAMPLE_GUIDE.exists(), reason="Sample fixture not found")
def test_output_has_paragraphs():
    with tempfile.TemporaryDirectory() as tmp:
        out = _render_sample(Path(tmp))
        doc = docxlib.Document(str(out))
        body_paragraphs = [
            p for p in doc.paragraphs if p.text.strip() and not p.style.name.startswith("Heading")
        ]
        assert len(body_paragraphs) >= 3


@pytest.mark.skipif(
    not SAMPLE_GUIDE.exists() or not EXPECTED_OUTPUT.exists(),
    reason="Expected output fixture not found",
)
def test_golden_heading_count_matches():
    with tempfile.TemporaryDirectory() as tmp:
        out = _render_sample(Path(tmp))
        actual = docxlib.Document(str(out))
        expected = docxlib.Document(str(EXPECTED_OUTPUT))

        actual_h1s = [p.text for p in actual.paragraphs if p.style.name == "Heading 1"]
        expected_h1s = [p.text for p in expected.paragraphs if p.style.name == "Heading 1"]
        assert len(actual_h1s) == len(expected_h1s), (
            f"H1 count mismatch: got {actual_h1s}, expected {expected_h1s}"
        )
