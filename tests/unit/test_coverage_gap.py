"""Tests targeting the remaining coverage gaps to push total above 80%."""

from __future__ import annotations

import json
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import httpx
import pytest
from fastapi import FastAPI
from httpx import AsyncClient

from docforge.server.auth import create_token
from docforge.server.routers import auth, documents, jobs, projects, system

# ---------------------------------------------------------------------------
# Test helpers
# ---------------------------------------------------------------------------

SECRET = "test-secret-32-chars-minimum-key!!"
TOKEN = create_token("admin", SECRET, ttl_hours=1)
AUTH = {"Authorization": f"Bearer {TOKEN}"}


def _make_app(tmp_path: Path) -> FastAPI:
    app = FastAPI()
    app.state.db_url = "postgresql://fake/fake"
    app.state.upload_dir = str(tmp_path / "uploads")
    app.state.output_dir = str(tmp_path / "output")
    app.state.secret_key = SECRET
    app.state.token_ttl_hours = 1
    (tmp_path / "uploads").mkdir()
    (tmp_path / "output").mkdir()
    app.include_router(auth.router)
    app.include_router(documents.router)
    app.include_router(jobs.router)
    app.include_router(projects.router)
    app.include_router(system.router)
    return app


# ---------------------------------------------------------------------------
# projects router — duplicate and delete with files
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_duplicate_project_not_found(tmp_path: Path) -> None:
    app = _make_app(tmp_path)
    with (
        patch("docforge.server.routers.projects.store.get_connection"),
        patch("docforge.server.routers.projects.store.get_project", return_value=None),
    ):
        async with AsyncClient(transport=httpx.ASGITransport(app=app), base_url="http://t") as c:
            r = await c.post("/projects/ghost/duplicate", headers=AUTH)
    assert r.status_code == 404


@pytest.mark.asyncio
async def test_duplicate_project_missing_input(tmp_path: Path) -> None:
    app = _make_app(tmp_path)
    project = {
        "id": "p1",
        "name": "T",
        "input_filename": "missing.docx",
        "config_snapshot": "{}",
        "output_paths": "[]",
        "template": "minimal",
        "language": "en",
        "ai_model": "gpt-4o",
    }
    with (
        patch("docforge.server.routers.projects.store.get_connection"),
        patch("docforge.server.routers.projects.store.get_project", return_value=project),
    ):
        async with AsyncClient(transport=httpx.ASGITransport(app=app), base_url="http://t") as c:
            r = await c.post("/projects/p1/duplicate", headers=AUTH)
    assert r.status_code == 409


@pytest.mark.asyncio
async def test_duplicate_project_success(tmp_path: Path) -> None:
    app = _make_app(tmp_path)
    # Create the input file so the existence check passes
    upload = tmp_path / "uploads"
    (upload / "doc.docx").write_bytes(b"fake")

    project = {
        "id": "p1",
        "name": "T",
        "input_filename": "doc.docx",
        "config_snapshot": "{}",
        "output_paths": "[]",
        "template": "minimal",
        "language": "en",
        "ai_model": "gpt-4o",
    }

    mock_queue = AsyncMock()
    mock_queue.submit = AsyncMock()

    with (
        patch("docforge.server.routers.projects.store.get_connection"),
        patch("docforge.server.routers.projects.store.get_project", return_value=project),
        patch("docforge.server.jobs.get_job_queue", return_value=mock_queue),
    ):
        async with AsyncClient(transport=httpx.ASGITransport(app=app), base_url="http://t") as c:
            r = await c.post("/projects/p1/duplicate", headers=AUTH)
    assert r.status_code == 202
    assert "job_id" in r.json()


@pytest.mark.asyncio
async def test_delete_project_removes_files(tmp_path: Path) -> None:
    app = _make_app(tmp_path)
    output_file = tmp_path / "output" / "job1.docx"
    output_file.write_bytes(b"data")

    project = {
        "id": "p1",
        "name": "T",
        "output_paths": json.dumps([str(output_file)]),
    }
    with (
        patch("docforge.server.routers.projects.store.get_connection"),
        patch("docforge.server.routers.projects.store.get_project", return_value=project),
        patch("docforge.server.routers.projects.store.delete_project", return_value=True),
    ):
        async with AsyncClient(transport=httpx.ASGITransport(app=app), base_url="http://t") as c:
            r = await c.delete("/projects/p1", headers=AUTH)
    assert r.status_code == 204
    assert not output_file.exists()


# ---------------------------------------------------------------------------
# layout validator
# ---------------------------------------------------------------------------


def test_layout_validator_heading_hierarchy_skip() -> None:
    from docforge.core.document import Chapter, SemanticModel
    from docforge.rendering.layout_validator import validate

    m = SemanticModel(document_id="d")
    m.chapters.append(Chapter(id="c1", title="C1", heading_level=1))
    m.chapters.append(Chapter(id="c3", title="C3", heading_level=3))  # skips level 2
    result = validate(m)
    codes = [w.code for w in result.warnings]
    assert "HEADING_HIERARCHY_SKIP" in codes


def test_layout_validator_orphan_heading() -> None:
    from docforge.core.document import Chapter, Heading, SemanticModel
    from docforge.rendering.layout_validator import validate

    m = SemanticModel(document_id="d")
    ch = Chapter(id="c1", title="C1", heading_level=1)
    ch.elements.append(Heading(text="A", level=2))
    ch.elements.append(Heading(text="B", level=3))
    m.chapters.append(ch)
    result = validate(m)
    codes = [w.code for w in result.warnings]
    assert "ORPHAN_HEADING" in codes


def test_layout_validator_oversized_placeholder() -> None:
    from docforge.core.document import Chapter, ImagePlaceholder, SemanticModel
    from docforge.rendering.layout_validator import validate

    m = SemanticModel(document_id="d")
    ch = Chapter(id="c1", title="C1", heading_level=1)
    ch.elements.append(ImagePlaceholder(placeholder_text="x" * 501))
    m.chapters.append(ch)
    result = validate(m)
    codes = [w.code for w in result.warnings]
    assert "OVERSIZED_PLACEHOLDER" in codes


def test_layout_validator_empty_table() -> None:
    from docforge.core.document import Chapter, SemanticModel, Table
    from docforge.rendering.layout_validator import validate

    m = SemanticModel(document_id="d")
    ch = Chapter(id="c1", title="C1", heading_level=1)
    ch.elements.append(Table(rows=[]))
    m.chapters.append(ch)
    result = validate(m)
    codes = [e.code for e in result.errors]
    assert "EMPTY_TABLE" in codes


def test_layout_validator_clean_model() -> None:
    from docforge.core.document import Chapter, Paragraph, SemanticModel
    from docforge.rendering.layout_validator import validate

    m = SemanticModel(document_id="d")
    ch = Chapter(id="c1", title="C1", heading_level=1)
    ch.elements.append(Paragraph(text="hello"))
    m.chapters.append(ch)
    result = validate(m)
    assert result.errors == []


# ---------------------------------------------------------------------------
# Renderer public API
# ---------------------------------------------------------------------------


def test_renderer_fluent_setters() -> None:
    from docforge.renderer import Renderer

    r = Renderer()
    r2 = r.template("minimal").language("ru").model("gpt-4o").creativity(7)
    assert r2 is r
    assert r._template == "minimal"
    assert r._language == "ru"
    assert r._model == "gpt-4o"
    assert r._creativity == 7


def test_renderer_creativity_invalid() -> None:
    from docforge.renderer import Renderer

    with pytest.raises(ValueError, match="creativity"):
        Renderer().creativity(0)

    with pytest.raises(ValueError, match="creativity"):
        Renderer().creativity(11)


def test_renderer_provider_setter() -> None:
    from docforge.renderer import Renderer

    fake = object()
    r = Renderer().provider(fake)
    assert r._ai_provider is fake


# ---------------------------------------------------------------------------
# document/loader.py — _extract_title and _detect_language fallbacks
# ---------------------------------------------------------------------------


def test_extract_title_from_core_properties(tmp_path: Path) -> None:
    from docforge.document.loader import _extract_title

    mock_doc = MagicMock()
    mock_doc.core_properties.title = "My Title"
    result = _extract_title(mock_doc, tmp_path / "doc.docx")
    assert result == "My Title"


def test_extract_title_from_heading(tmp_path: Path) -> None:
    from docforge.document.loader import _extract_title

    para = MagicMock()
    para.style.name = "Heading 1"
    para.text = "  Big Title  "

    mock_doc = MagicMock()
    mock_doc.core_properties.title = ""
    mock_doc.paragraphs = [para]
    result = _extract_title(mock_doc, tmp_path / "doc.docx")
    assert result == "Big Title"


def test_extract_title_falls_back_to_stem(tmp_path: Path) -> None:
    from docforge.document.loader import _extract_title

    para = MagicMock()
    para.style.name = "Normal"
    para.text = "body text"

    mock_doc = MagicMock()
    mock_doc.core_properties.title = ""
    mock_doc.paragraphs = [para]
    result = _extract_title(mock_doc, tmp_path / "my_document.docx")
    assert result == "my_document"


def test_detect_language_fallback_to_en(tmp_path: Path) -> None:
    from docforge.document.loader import _detect_language

    mock_doc = MagicMock()
    # Simulate exception in body traversal
    mock_doc.element.body.iter.side_effect = Exception("boom")
    result = _detect_language(mock_doc)
    assert result == "en"


# ---------------------------------------------------------------------------
# Plugin registry — load_from_directory
# ---------------------------------------------------------------------------


def test_registry_load_from_directory_skips_bad_manifests(tmp_path: Path) -> None:
    from docforge.plugins.registry import PluginRegistry

    # Create a directory with a malformed plugin.yaml
    plugin_dir = tmp_path / "bad_plugin"
    plugin_dir.mkdir()
    (plugin_dir / "plugin.yaml").write_text("id: bad\nentrypoint: nonexistent.module\n")

    reg = PluginRegistry()
    reg.load_from_directory(tmp_path)
    # Didn't raise; bad plugin was skipped
    assert reg.get("bad") is None


def test_registry_load_from_directory_empty(tmp_path: Path) -> None:
    from docforge.plugins.registry import PluginRegistry

    reg = PluginRegistry()
    reg.load_from_directory(tmp_path)  # no plugin.yaml files
    assert reg.all == {}


# ---------------------------------------------------------------------------
# RenderingReport
# ---------------------------------------------------------------------------


def test_rendering_report_full_lifecycle() -> None:
    from docforge.rendering.report import RenderingReport

    r = RenderingReport(job_id="j1", input_filename="doc.docx")
    r.add_warning("w1")
    r.add_recovered_error("e1")
    r.add_skipped("s1")
    assert r.warnings == ["w1"]
    assert r.recovered_errors == ["e1"]
    assert r.skipped_operations == ["s1"]
    assert r.succeeded()

    r.finish()
    assert r.completed_at is not None
    assert r.duration_seconds >= 0

    s = r.summary()
    assert s["job_id"] == "j1"
    assert s["succeeded"] is True
    assert s["warnings"] == 1


def test_rendering_report_failure() -> None:
    from docforge.rendering.report import RenderingReport

    r = RenderingReport(job_id="j2", input_filename="doc.docx")
    r.fatal_failure = "boom"
    assert not r.succeeded()
    assert r.summary()["fatal"] == "boom"


def test_rendering_report_image_attribution() -> None:
    from docforge.rendering.report import RenderingReport

    r = RenderingReport(job_id="j3", input_filename="doc.docx")

    class FakeCandidate:
        provider = "wikimedia"
        title = "Cat"
        author = "Alice"
        url = "https://example.com/cat.jpg"
        source_page = "https://example.com"
        licence = "CC-BY"

    r.add_image_attribution(FakeCandidate())
    assert len(r.image_attributions) == 1
    assert r.image_attributions[0]["title"] == "Cat"
    assert r.image_attributions[0]["author"] == "Alice"


# ---------------------------------------------------------------------------
# EPUB exporter via mock (no ebooklib installed in CI)
# ---------------------------------------------------------------------------


def test_epub_export_with_mock(tmp_path: Path) -> None:
    """Test epub exporter code paths using a fully mocked ebooklib."""
    from unittest.mock import MagicMock, patch

    from docforge.core.document import Chapter, Paragraph, SemanticModel

    m = SemanticModel(document_id="d")
    ch = Chapter(id="c0", title="Ch1", heading_level=1)
    ch.elements.append(Paragraph(text="Hello."))
    m.chapters.append(ch)

    mock_epub_book = MagicMock()
    mock_epub_html = MagicMock()
    mock_epub_html.file_name = "chapter_000.xhtml"
    mock_epub_html.id = "ch-id"
    mock_epub_html.title = "Ch1"
    mock_epub_item = MagicMock()

    mock_epub_mod = MagicMock()
    mock_epub_mod.EpubBook.return_value = mock_epub_book
    mock_epub_mod.EpubHtml.return_value = mock_epub_html
    mock_epub_mod.EpubItem.return_value = mock_epub_item
    mock_epub_mod.EpubNcx.return_value = MagicMock()
    mock_epub_mod.EpubNav.return_value = MagicMock()
    mock_epub_mod.Link.return_value = MagicMock()
    mock_epub_mod.write_epub = MagicMock()

    out = tmp_path / "out.epub"
    # Simulate write_epub creating the file
    mock_epub_mod.write_epub.side_effect = lambda path, book: Path(path).write_bytes(b"epub")

    import sys

    mock_ebooklib = MagicMock()
    mock_ebooklib.epub = mock_epub_mod

    with patch.dict(sys.modules, {"ebooklib": mock_ebooklib, "ebooklib.epub": mock_epub_mod}):
        import importlib

        import docforge.exporters.epub as epub_mod

        importlib.reload(epub_mod)
        result = epub_mod.export(m, out, title="Test Book", language="en")

    assert result == out


# ---------------------------------------------------------------------------
# jobs router — estimate endpoint
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_estimate_job_with_config(tmp_path: Path) -> None:
    app = _make_app(tmp_path)

    docx = tmp_path / "uploads" / "doc1.docx"
    # Create a real minimal docx so analyse() can parse it
    import io

    from docx import Document as DocxDocument

    buf = io.BytesIO()
    d = DocxDocument()
    d.add_heading("Chapter One", level=1)
    d.add_paragraph("Some text here.")
    d.save(buf)
    docx.write_bytes(buf.getvalue())

    async with AsyncClient(transport=httpx.ASGITransport(app=app), base_url="http://t") as c:
        r = await c.post(
            "/jobs/estimate",
            json={
                "document_id": "doc1",
                "config": {
                    "coverPage": "none",
                    "tableOfContents": "keep_existing",
                    "headersFooters": "keep_existing",
                },
            },
            headers=AUTH,
        )
    assert r.status_code == 200
    data = r.json()
    assert "estimated_ai_cost_usd" in data
    assert data["has_cover_page"] is False
    assert data["has_toc"] is False
    assert data["has_headers_footers"] is False


@pytest.mark.asyncio
async def test_estimate_job_default_config(tmp_path: Path) -> None:
    app = _make_app(tmp_path)
    docx = tmp_path / "uploads" / "doc2.docx"
    import io

    from docx import Document as DocxDocument

    buf = io.BytesIO()
    d = DocxDocument()
    d.add_heading("Chapter One", level=1)
    d.add_paragraph("Some text here.")
    d.save(buf)
    docx.write_bytes(buf.getvalue())

    async with AsyncClient(transport=httpx.ASGITransport(app=app), base_url="http://t") as c:
        r = await c.post(
            "/jobs/estimate",
            json={"document_id": "doc2"},
            headers=AUTH,
        )
    assert r.status_code == 200
    data = r.json()
    assert data["has_cover_page"] is True
    assert data["has_toc"] is True
